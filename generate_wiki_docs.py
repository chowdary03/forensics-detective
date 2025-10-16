#!/usr/bin/env python3
"""
Scrape real Wikipedia articles and generate DOCX documents
for ForensicsDetective project.

- Pulls live Wikipedia content from diverse domains
- Adds random valid images
- Appends new docs safely to wikipedia_docs/
"""

import os
import random
from docx import Document
from docx.shared import Inches
from tqdm import tqdm
from PIL import Image as PILImage
import wikipedia

# === Folder paths ===
OUTPUT_FOLDER = "wikipedia_docs"
IMAGES_FOLDER = "sample_images"
NUM_DOCS = 5000  # total to generate

# === Ensure output folder exists ===
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# === Topic pools (diverse Wikipedia coverage) ===
TOPICS = {
    "STEM": [
        "Quantum mechanics", "Relativity", "Machine learning", "Data structures",
        "Artificial intelligence", "Computer vision", "Linear algebra", "Blockchain",
        "Neural networks", "Genetic engineering"
    ],
    "Humanities": [
        "Renaissance art", "Greek mythology", "Shakespeare", "Philosophy of mind",
        "Ethics", "World literature", "History of music", "Feminist theory"
    ],
    "Social Sciences": [
        "Cognitive psychology", "Sociology", "Political economy",
        "Cultural anthropology", "Economics", "Education", "Social stratification"
    ],
    "Contemporary": [
        "Climate change", "Renewable energy", "Artificial intelligence ethics",
        "Globalization", "Cybersecurity", "Social media impact"
    ],
    "Reference": [
        "Python programming", "How to write a research paper",
        "Technical documentation", "Scientific method", "Open source software"
    ]
}

# === Function to add a random valid image ===
def add_random_image(doc):
    """Add a random valid image from IMAGES_FOLDER."""
    if not os.path.exists(IMAGES_FOLDER):
        return
    valid_exts = [".jpg", ".jpeg", ".png", ".bmp", ".gif"]
    images = [f for f in os.listdir(IMAGES_FOLDER)
              if os.path.splitext(f)[1].lower() in valid_exts]
    if not images:
        return
    random.shuffle(images)
    for img_file in images:
        img_path = os.path.join(IMAGES_FOLDER, img_file)
        try:
            with PILImage.open(img_path) as im:
                im.verify()
            doc.add_picture(img_path, width=Inches(random.uniform(2, 4)))
            doc.add_paragraph("")  # spacing
            return
        except Exception:
            continue

# === Determine starting index based on existing files ===
existing_files = [f for f in os.listdir(OUTPUT_FOLDER) if f.endswith(".docx")]
existing_indices = [
    int(f.split('_')[-1].split('.')[0])
    for f in existing_files if f.startswith("wiki_doc_")
]
start_index = max(existing_indices, default=0) + 1

print(f"Generating {NUM_DOCS} real Wikipedia documents starting from index {start_index} ...")

# === Main generation loop ===
for i in tqdm(range(start_index, start_index + NUM_DOCS)):
    # Choose a random topic from all categories
    category = random.choice(list(TOPICS.keys()))
    topic = random.choice(TOPICS[category])
    doc = Document()
    doc.add_heading(topic, level=0)

    try:
        # Fetch Wikipedia content
        page = wikipedia.page(topic, auto_suggest=False)
        content = page.content

        # Split long content into smaller sections
        paragraphs = content.split("\n")
        for para in paragraphs[:15]:  # limit paragraphs per doc
            if para.strip():
                doc.add_paragraph(para.strip())

        # Occasionally add image and table
        if random.random() < 0.3:
            add_random_image(doc)
        if random.random() < 0.25:
            table = doc.add_table(rows=3, cols=3)
            for r in table.rows:
                for c in r.cells:
                    c.text = str(random.randint(0, 100))

        # Save file
        output_path = os.path.join(OUTPUT_FOLDER, f"wiki_doc_{i:05d}.docx")
        doc.save(output_path)

    except Exception as e:
        print(f"⚠️ Skipping topic '{topic}' due to error: {e}")

print(f"\n✅ Completed generating {NUM_DOCS} Wikipedia-based documents in '{OUTPUT_FOLDER}/'")
