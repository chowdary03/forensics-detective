import random
import os
from docx import Document
from docx.shared import Inches

folder = "wikipedia_docs_scaled"
image_folder = "sample_images"  # folder with images

# List all image files
image_files = [f for f in os.listdir(image_folder) if f.endswith((".png", ".jpg", ".jpeg"))]

# Pick 3000 random files to add images
files = [f for f in os.listdir(folder) if f.endswith(".docx")]
image_docs = random.sample(files, 3000)

for file in image_docs:
    path = os.path.join(folder, file)
    doc = Document(path)
    
    # Add heading
    doc.add_heading("Inserted Image Section", level=2)
    
    # Pick a random image from sample_images
    img = random.choice(image_files)
    img_path = os.path.join(image_folder, img)
    
    # Insert image
    doc.add_picture(img_path, width=Inches(2))
    
    doc.save(path)
